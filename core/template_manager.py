"""
Работа с шаблонами .docx и metadata (.json).

На основе спецификации:
- scan_templates(folder)           — поиск шаблонов и метаданных;
- load_template_metadata(path)     — загрузка template.json;
- extract_keys_from_docx(path)     — извлечение переменных из docx (docxtpl).

Сейчас реализован каркас без фактического использования docxtpl,
чтобы приложение можно было запускать без дополнительных зависимостей.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any
import re

from docx import Document


logger = logging.getLogger(__name__)


TEMPLATES_DIR = Path("templates")


DOC_SPECIFIC_KEYWORDS = [
    "материал",
    "серт",
    "прилож",
    "схем",
    "черт",
    "испыт",
    "наименование_работ",
]

# Явные поля из рабочего сценария (см. присланный пример Nik66.py).
KNOWN_CONSTANT_KEYS = {
    "Объект",
    "Застройщик",
    "Строитель",
    "Проектная_организация",
    "Проект_или_ТЗ",
    "Представитель_застр",
    "ФИО_застр",
    "Распор_застр",
    "Пр_раб",
    "ФИО_Пр_раб",
    "Распор_пр_раб",
    "Строй_контроль_Должность",
    "ФИО_Стройк",
    "Распор_стройк",
    "Проектировщик_должность",
    "Проектировщик_ФИО",
    "Распоряжение_проект",
    "Выполнил_работы",
    "Иные_долж",
    "ФИО_Иные",
    "Распор_иные",
    "Организация_исполнитель",
    "Экз",
}

KNOWN_DOC_KEYS = {
    "номер",
    "Наименование_работ",
    "Начало",
    "Конец",
    "Материалы_и_серты",
    "Схемы_и_тд",
    "Разрешает_пр_во_работ_по",
    "СП",
    "Ч",
    "М",
    "Г",
    "Чнач",
    "Мнач",
    "Гн",
    "date_end",
}


def split_fields_by_scope(keys: List[str]) -> tuple[List[str], List[str]]:
    """Разделяет ключи на постоянные и по-документные с сохранением порядка.

    Приоритет:
    1) явные списки KNOWN_CONSTANT_KEYS / KNOWN_DOC_KEYS;
    2) fallback-эвристика по ключевым словам для по-актовых полей.
    """
    constant_keys: List[str] = []
    doc_keys: List[str] = []

    for key in keys:
        normalized = (key or "").strip()
        if not normalized:
            continue

        if normalized in KNOWN_CONSTANT_KEYS:
            constant_keys.append(normalized)
            continue

        if normalized in KNOWN_DOC_KEYS:
            doc_keys.append(normalized)
            continue

        low = normalized.lower()
        is_doc_specific = any(word in low for word in DOC_SPECIFIC_KEYWORDS)
        if is_doc_specific:
            doc_keys.append(normalized)
        else:
            constant_keys.append(normalized)

    return constant_keys, doc_keys



@dataclass
class TemplateMeta:
    """Метаданные одного шаблона .docx."""

    path: Path
    display_name: str
    type: str
    multiple: bool = False
    prefix: Optional[str] = None
    date_mode: str = "single"  # "single" | "period"
    in_registry: bool = True
    default_pages: int = 1
    fields: Dict[str, str] | None = None
    constant_fields_order: List[str] | None = None
    doc_fields_order: List[str] | None = None

    def to_dict(self) -> Dict[str, Any]:
        """Сериализация в dict формата template.json (без абсолютных путей)."""
        return {
            "display_name": self.display_name,
            "type": self.type,
            "multiple": self.multiple,
            "prefix": self.prefix,
            "date_mode": self.date_mode,
            "in_registry": self.in_registry,
            "default_pages": self.default_pages,
            "fields": self.fields or {},
            "constant_fields_order": self.constant_fields_order or [],
            "doc_fields_order": self.doc_fields_order or [],
        }


def scan_templates(folder: str | Path | None = None) -> List[TemplateMeta]:
    """Сканирует папку с шаблонами и ищет рядом JSON‑метаданные.

    Ожидается структура:
      templates/
        01_AOSR.docx
        01_AOSR.json
    """
    base = Path(folder) if folder is not None else TEMPLATES_DIR
    templates: List[TemplateMeta] = []

    for docx_path in sorted(base.glob("*.docx")):
        meta_path = docx_path.with_suffix(".json")
        meta = load_template_metadata(meta_path) if meta_path.exists() else None
        if meta is None:
            # Минимальные метаданные по умолчанию
            templates.append(
                TemplateMeta(
                    path=docx_path,
                    display_name=docx_path.stem,
                    type="GENERIC",
                )
            )
        else:
            templates.append(meta)

    return templates


def load_template_metadata(path: str | Path) -> TemplateMeta | None:
    """Загружает metadata для шаблона из JSON‑файла.

    Если формат неожиданен, возвращает None.
    """
    p = Path(path)
    with p.open("r", encoding="utf-8") as f:
        data = json.load(f)

    try:
        return TemplateMeta(
            path=p.with_suffix(".docx"),
            display_name=data.get("display_name") or p.stem,
            type=data.get("type") or "GENERIC",
            multiple=bool(data.get("multiple", False)),
            prefix=data.get("prefix"),
            date_mode=data.get("date_mode", "single"),
            in_registry=bool(data.get("in_registry", True)),
            default_pages=int(data.get("default_pages", 1)),
            fields=data.get("fields") or {},
            constant_fields_order=data.get("constant_fields_order") or [],
            doc_fields_order=data.get("doc_fields_order") or [],
        )
    except Exception as exc:
        logger.error("Failed to parse template metadata %s: %s", p, exc)
        return None


def load_template_metadata_from_dict(docx_path: str | Path, data: Dict[str, Any]) -> TemplateMeta | None:
    """Создаёт TemplateMeta из словаря metadata (как в template.json)."""
    try:
        p = Path(docx_path)
        return TemplateMeta(
            path=p,
            display_name=data.get("display_name") or p.stem,
            type=data.get("type") or "GENERIC",
            multiple=bool(data.get("multiple", False)),
            prefix=data.get("prefix"),
            date_mode=data.get("date_mode", "single"),
            in_registry=bool(data.get("in_registry", True)),
            default_pages=int(data.get("default_pages", 1)),
            fields=data.get("fields") or {},
            constant_fields_order=data.get("constant_fields_order") or [],
            doc_fields_order=data.get("doc_fields_order") or [],
        )
    except Exception as exc:
        logger.error("Failed to build TemplateMeta from dict for %s: %s", docx_path, exc)
        return None


def extract_keys_from_docx(path: str | Path) -> List[str]:
    """Извлекает плейсхолдеры {{ключ}} в порядке следования в документе.

    Для сохранения порядка читаем текст документа и не используем set как итоговую структуру.
    Порядок ключей используется затем для полей metadata и
    для добавления ключей в "Постоянные данные".
    """
    p = Path(path)
    try:
        doc = Document(p)
    except Exception as exc:
        logger.error("Failed to open docx %s: %s", p, exc)
        return []

    chunks: List[str] = []

    # Тело документа
    chunks.extend(par.text for par in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(par.text for par in cell.paragraphs)

    # Колонтитулы и заголовки секций
    for section in doc.sections:
        chunks.extend(par.text for par in section.header.paragraphs)
        chunks.extend(par.text for par in section.footer.paragraphs)

    text = "\n".join(chunks)

    pattern = re.compile(r"{{\s*([^{}\n]+?)\s*}}")
    seen = set()
    keys: List[str] = []
    for m in pattern.finditer(text):
        key = m.group(1)
        if key not in seen:
            seen.add(key)
            keys.append(key)
    return keys


def build_metadata_skeleton(docx_path: str | Path) -> TemplateMeta:
    """Строит объект TemplateMeta по одному .docx, извлекая ключи.

    Полезно для автогенерации начального template.json.
    Тип по умолчанию — GENERIC, поля из ключей помечаются типом 'string'.
    """
    p = Path(docx_path)
    keys = extract_keys_from_docx(p)
    # Важно: не сортируем keys, чтобы сохранить порядок,
    # максимально близкий к порядку появления в шаблоне.
    fields: Dict[str, str] = {k: "string" for k in keys}
    constant_fields_order, doc_fields_order = split_fields_by_scope(keys)

    return TemplateMeta(
        path=p,
        display_name=p.stem,
        type="GENERIC",
        multiple=False,
        prefix=None,
        date_mode="single",
        in_registry=True,
        default_pages=1,
        fields=fields,
        constant_fields_order=constant_fields_order,
        doc_fields_order=doc_fields_order,
    )


def save_template_metadata(meta: TemplateMeta, json_path: str | Path | None = None) -> Path:
    """Сохраняет TemplateMeta в JSON рядом с .docx и возвращает путь к файлу."""
    if json_path is None:
        json_path = meta.path.with_suffix(".json")
    p = Path(json_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    with p.open("w", encoding="utf-8") as f:
        json.dump(meta.to_dict(), f, ensure_ascii=False, indent=2)
    return p
